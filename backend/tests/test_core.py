"""Юнит-тесты llm_client, scoring, parser."""

from __future__ import annotations

import json
from unittest.mock import MagicMock, patch

import pytest

from app.ingest.ocr import pdf_needs_ocr, tesseract_available
from app.ingest.parser import make_doc_id, parse_entities_from_llm, _read_docx
from app.ingest.text_utils import chunk_text, text_quality_score
from app.llm_client import YandexLLMClient
from app.models import Hypothesis, RiskScores, ScoreBreakdown
from app.scoring.ranker import Ranker


class TestParser:
    def test_chunk_text_splits_long_text(self):
        text = "word " * 500
        chunks = chunk_text(text, chunk_size=100, overlap=20)
        assert len(chunks) > 1
        assert all(len(c) <= 120 for c in chunks)

    def test_chunk_text_short(self):
        assert chunk_text("short text") == ["short text"]

    def test_make_doc_id_stable(self):
        id1 = make_doc_id(__import__("pathlib").Path("test.pdf"), "content")
        id2 = make_doc_id(__import__("pathlib").Path("test.pdf"), "content")
        assert id1 == id2

    def test_parse_entities_from_llm(self):
        data = {
            "entities": [
                {"name": "Cu", "type": "Material", "properties": {}},
                {"name": "", "type": "Process"},
            ]
        }
        entities = parse_entities_from_llm(data)
        assert len(entities) == 1
        assert entities[0].name == "Cu"

    def test_pdf_needs_ocr_heuristics(self):
        assert pdf_needs_ocr("", 10) is True
        assert pdf_needs_ocr("x" * 50, 5) is True
        assert pdf_needs_ocr("x" * 5000, 10) is False

    def test_read_docx_includes_tables(self, tmp_path):
        from docx import Document

        path = tmp_path / "tables.docx"
        doc = Document()
        doc.add_paragraph("Вводный текст")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Медь"
        table.cell(0, 1).text = "95%"
        table.cell(1, 0).text = "Цинк"
        table.cell(1, 1).text = "12%"
        doc.save(str(path))

        text = _read_docx(path)
        assert "Вводный текст" in text
        assert "Медь" in text
        assert "95%" in text
        assert "## Table 1" in text

    def test_clean_ocr_filters_noise(self):
        from app.ingest.text_utils import clean_ocr_text, text_quality_score

        noisy = "||| @@ ## \nфлотация меди\n||| @@"
        cleaned = clean_ocr_text(noisy)
        assert "флотация" in cleaned
        assert text_quality_score("Нормальный текст о флотации медных руд") > 0.3

    def test_smart_chunk_preserves_paragraphs(self):
        text = "Абзац один. " * 50 + "\n\n" + "Абзац два. " * 50
        chunks = chunk_text(text, chunk_size=200, overlap=30)
        assert len(chunks) >= 2


class TestScoring:
    def _make_hypothesis(self, **kwargs) -> Hypothesis:
        defaults = {
            "id": "h1",
            "text": "Test hypothesis",
            "mechanism": "mechanism",
            "novelty_score": 8.0,
            "feasibility_score": 6.0,
            "expected_value_score": 7.0,
            "risk": RiskScores(technical=3.0, economic=4.0),
        }
        defaults.update(kwargs)
        return Hypothesis(**defaults)

    def test_ranker_composite_score(self):
        ranker = Ranker(llm=None, qdrant=None)
        h = ranker.score_hypothesis(self._make_hypothesis())
        assert h.score_breakdown is not None
        assert 0 <= h.score_breakdown.composite <= 1

    def test_ranker_sorts_descending(self):
        ranker = Ranker(llm=None, qdrant=None)
        h1 = self._make_hypothesis(id="h1", novelty_score=9)
        h2 = self._make_hypothesis(id="h2", novelty_score=3)
        ranked = ranker.rank([h2, h1])
        assert ranked[0].id == "h1"

    def test_custom_weights(self):
        ranker = Ranker(llm=None, qdrant=None)
        weights = {"novelty": 0.9, "feasibility": 0.05, "expected_value": 0.03, "risk": 0.02}
        h = ranker.score_hypothesis(self._make_hypothesis(), weights=weights)
        assert h.score_breakdown.weights["novelty"] == 0.9


class TestLLMClient:
    def test_parse_json_with_markdown_fence(self):
        raw = '```json\n{"hypotheses": []}\n```'
        result = YandexLLMClient._parse_json(raw)
        assert result == {"hypotheses": []}

    def test_parse_json_embedded(self):
        raw = 'Some text {"key": "value"} trailing'
        result = YandexLLMClient._parse_json(raw)
        assert result["key"] == "value"

    def test_parse_json_control_chars_in_string(self):
        raw = '{"text": "line1\nline2"}'
        result = YandexLLMClient._parse_json(raw)
        assert result["text"] == "line1\nline2"

    def test_parse_json_rejects_null(self):
        import pytest

        with pytest.raises(ValueError, match="JSON-объекта"):
            YandexLLMClient._parse_json("null")

    def test_init_requires_credentials(self):
        with patch("app.llm_client.settings") as mock_settings:
            mock_settings.yc_api_key = ""
            mock_settings.yc_folder_id = ""
            with pytest.raises(ValueError, match="YC_API_KEY"):
                YandexLLMClient()

    @patch("app.llm_client.httpx.Client")
    def test_complete_success(self, mock_client_cls):
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.json.return_value = {
            "result": {
                "alternatives": [{"message": {"text": "ответ"}}],
                "usage": {"inputTextTokens": 10, "completionTokens": 5, "totalTokens": 15},
            }
        }
        mock_client = MagicMock()
        mock_client.__enter__ = MagicMock(return_value=mock_client)
        mock_client.__exit__ = MagicMock(return_value=False)
        mock_client.post.return_value = mock_response
        mock_client_cls.return_value = mock_client

        with patch("app.llm_client.settings") as mock_settings:
            mock_settings.yc_api_key = "test-key"
            mock_settings.yc_folder_id = "folder"
            mock_settings.yandexgpt_model = "yandexgpt"
            mock_settings.llm_completion_url = "http://test/completion"
            mock_settings.llm_temperature = 0.3
            mock_settings.llm_max_tokens = 1000
            mock_settings.llm_request_delay_sec = 0

            client = YandexLLMClient()
            result = client.complete("system", "user")
            assert result == "ответ"
            assert len(client.call_history) == 1
