"""Экспорт отчётов PDF/DOCX."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models import Hypothesis


def _styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleRu", parent=styles["Title"], fontSize=16, spaceAfter=12))
    styles.add(ParagraphStyle(name="BodyRu", parent=styles["Normal"], fontSize=10, leading=14))
    return styles


def export_pdf(
    problem: str,
    constraints: str,
    hypotheses: list[Hypothesis],
    output_path: Path,
) -> Path:
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm)
    styles = _styles()
    story: list[Any] = []

    story.append(Paragraph("Отчёт: научные гипотезы", styles["TitleRu"]))
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    story.append(Paragraph(f"Дата: {now}", styles["BodyRu"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Проблема:</b> {problem}", styles["BodyRu"]))
    story.append(Paragraph(f"<b>Ограничения:</b> {constraints or '—'}", styles["BodyRu"]))
    story.append(Spacer(1, 16))

    for i, h in enumerate(hypotheses, 1):
        score = h.score_breakdown.composite if h.score_breakdown else 0
        story.append(Paragraph(f"<b>{i}. {h.text}</b> (score: {score:.3f})", styles["BodyRu"]))
        story.append(Paragraph(f"Механизм: {h.mechanism}", styles["BodyRu"]))
        story.append(
            Paragraph(
                f"Новизна={h.novelty_score}, Реализуемость={h.feasibility_score}, "
                f"Ценность={h.expected_value_score}, Риск тех/экон={h.risk.technical}/{h.risk.economic}",
                styles["BodyRu"],
            )
        )
        if h.reasoning:
            story.append(Paragraph(f"Обоснование: {h.reasoning[:500]}", styles["BodyRu"]))
        if h.business_case and h.business_case.narrative:
            bc = h.business_case
            story.append(
                Paragraph(
                    f"Бизнес-кейс: KPI={bc.target_kpi}, "
                    f"Δ={bc.expected_delta_pct or '—'}%, "
                    f"ROI={bc.roi_ratio or '—'}",
                    styles["BodyRu"],
                )
            )
        if h.sources:
            src = "; ".join(f"{s.doc_id}: {s.snippet[:80]}" for s in h.sources[:3])
            story.append(Paragraph(f"Источники: {src}", styles["BodyRu"]))
        story.append(Spacer(1, 10))

    doc.build(story)
    return output_path


def export_docx(
    problem: str,
    constraints: str,
    hypotheses: list[Hypothesis],
    output_path: Path,
) -> Path:
    from docx import Document

    document = Document()
    document.add_heading("Отчёт: научные гипотезы", 0)
    document.add_paragraph(f"Дата: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    document.add_paragraph(f"Проблема: {problem}")
    document.add_paragraph(f"Ограничения: {constraints or '—'}")

    for i, h in enumerate(hypotheses, 1):
        score = h.score_breakdown.composite if h.score_breakdown else 0
        document.add_heading(f"{i}. {h.text} (score: {score:.3f})", level=2)
        document.add_paragraph(f"Механизм: {h.mechanism}")
        document.add_paragraph(
            f"Новизна={h.novelty_score}, Реализуемость={h.feasibility_score}, "
            f"Ценность={h.expected_value_score}"
        )
        if h.reasoning:
            document.add_paragraph(h.reasoning)
        if h.verification_roadmap:
            document.add_paragraph("Дорожная карта верификации:")
            for step in h.verification_roadmap:
                document.add_paragraph(step, style="List Bullet")
        if h.sources:
            document.add_paragraph("Источники:")
            for s in h.sources:
                document.add_paragraph(f"{s.doc_id}: {s.snippet[:200]}", style="List Bullet")

    document.save(str(output_path))
    return output_path


def export_report(
    problem: str,
    constraints: str,
    hypotheses: list[Hypothesis],
    fmt: str,
    output_dir: Path,
    generation_id: str,
) -> tuple[Path, bytes]:
    output_dir.mkdir(parents=True, exist_ok=True)
    if fmt == "docx":
        path = output_dir / f"report_{generation_id[:8]}.docx"
        export_docx(problem, constraints, hypotheses, path)
    else:
        path = output_dir / f"report_{generation_id[:8]}.pdf"
        export_pdf(problem, constraints, hypotheses, path)
    return path, path.read_bytes()
