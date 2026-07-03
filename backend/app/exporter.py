"""Экспорт: DOCX-отчёт, CSV/JSON задач для Jira/YouTrack."""
from __future__ import annotations

import csv
import io
import json
from datetime import date

from docx import Document
from docx.shared import Pt


def report_docx(run: dict) -> bytes:
    """Бизнес-отчёт: диагноз + ранжированные гипотезы с обоснованием и источниками."""
    doc = Document()
    doc.add_heading("Отчёт: гипотезы по снижению потерь металлов с хвостами", level=0)
    doc.add_paragraph(f"Файл данных: {run.get('input_file', '—')}  |  Дата: {date.today().isoformat()}")
    doc.add_paragraph(f"Цель: {run.get('goal') or '—'}")
    doc.add_paragraph(f"Ограничения: {run.get('constraints') or '—'}")

    doc.add_heading("Диагноз потерь", level=1)
    for line in (run.get("summary_text") or "").splitlines():
        p = doc.add_paragraph(line.strip("# "))
        p.paragraph_format.space_after = Pt(2)

    doc.add_heading("Гипотезы (по убыванию приоритета)", level=1)
    for i, h in enumerate(run.get("hypotheses", []), 1):
        r = h.get("ranking", {})
        doc.add_heading(f"{i}. {h['hypothesis']}", level=2)
        doc.add_paragraph(f"Итоговый скор: {r.get('final', '—')}   "
                          f"(новизна/реализуемость/эффект/риск: "
                          + "/".join(str((h.get('scores') or {}).get(k, '—'))
                                     for k in ("novelty", "feasibility", "impact", "risk")) + ")")
        doc.add_paragraph(f"Механизм: {h.get('mechanism', '—')}")
        doc.add_paragraph(f"Ожидаемый эффект: {h.get('expected_effect', '—')}")
        doc.add_paragraph(f"Реализуемость: {h.get('feasibility_note', '—')}")
        risks = h.get("risks") or {}
        doc.add_paragraph(f"Риски — технические: {risks.get('technical', '—')}; "
                          f"экономические: {risks.get('economic', '—')}")
        if h.get("verification_roadmap"):
            doc.add_paragraph("Дорожная карта проверки:")
            for step in h["verification_roadmap"]:
                doc.add_paragraph(step, style="List Bullet")
        if h.get("sources"):
            doc.add_paragraph("Источники:")
            for s in h["sources"]:
                doc.add_paragraph(f"[{s['ref']}] {s['doc_id']}, стр. {s['page']}: «{s['snippet'][:200]}…»",
                                  style="List Bullet")
        if not h.get("grounded"):
            doc.add_paragraph("⚠ Гипотеза не подтверждена источниками базы знаний — требует экспертной проверки.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def tasks_csv(run: dict) -> str:
    buf = io.StringIO()
    w = csv.writer(buf, delimiter=";")
    w.writerow(["Summary", "Description", "Priority", "Score", "Element", "SizeClass", "Sources"])
    for i, h in enumerate(run.get("hypotheses", []), 1):
        t = h.get("target") or {}
        srcs = "; ".join(f"{s['doc_id']} стр.{s['page']}" for s in h.get("sources", []))
        w.writerow([
            f"Проверить гипотезу: {h['hypothesis'][:120]}",
            f"{h.get('mechanism', '')}\nЭффект: {h.get('expected_effect', '')}\n"
            f"Дорожная карта: {' | '.join(h.get('verification_roadmap', []))}",
            "High" if i <= 3 else "Medium",
            h.get("ranking", {}).get("final", ""),
            t.get("element", ""),
            t.get("size_class", ""),
            srcs,
        ])
    return buf.getvalue()


def tasks_json(run: dict) -> str:
    tasks = []
    for i, h in enumerate(run.get("hypotheses", []), 1):
        tasks.append({
            "summary": f"Проверить гипотезу: {h['hypothesis'][:120]}",
            "description": h.get("mechanism", ""),
            "expected_effect": h.get("expected_effect", ""),
            "priority": "High" if i <= 3 else "Medium",
            "score": h.get("ranking", {}).get("final"),
            "score_breakdown": h.get("ranking", {}).get("components"),
            "target": h.get("target"),
            "risks": h.get("risks"),
            "sources": h.get("sources"),
            "verification_roadmap": h.get("verification_roadmap"),
        })
    return json.dumps(tasks, ensure_ascii=False, indent=2)
