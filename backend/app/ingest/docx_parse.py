"""Универсальный парсинг DOCX — списки, таблицы, нумерованные пункты."""

from __future__ import annotations

import re
from pathlib import Path


def _numbered_line(line: str) -> tuple[str, str] | None:
    """Извлекает номер и текст из строки: '1. ...', '2) ...', '| 1. ... |'."""
    stripped = line.strip()
    m = re.match(r"^(?:\d+\.\s*|\d+\)\s*)(.+)$", stripped)
    if m:
        num = re.match(r"^(\d+)", stripped)
        return (num.group(1), m.group(1).strip()) if num else None

    if "|" in stripped:
        cells = [c.strip() for c in stripped.split("|") if c.strip()]
        for cell in cells:
            m2 = re.match(r"^(\d+)\.\s*(.+)$", cell)
            if m2:
                return m2.group(1), m2.group(2).strip()
    return None


def _has_numbered_block(text: str, *, min_items: int = 2) -> bool:
    count = sum(1 for line in text.splitlines() if _numbered_line(line))
    return count >= min_items


def enrich_structured_docx(text: str, path: Path) -> str:
    """
    Обогащает docx с нумерованными пунктами (гипотезы, рекомендации, шаги).
    Срабатывает по содержимому, не только по имени файла.
    """
    is_brainstorm = "мозгового штурма" in text.lower()
    is_hypotheses_file = "гипотез" in path.name.lower()
    has_numbered = _has_numbered_block(text)

    if not (is_brainstorm or is_hypotheses_file or has_numbered):
        return text

    topics: list[str] = []
    enriched_lines: list[str] = []
    for line in text.splitlines():
        parsed = _numbered_line(line)
        if parsed:
            num, topic = parsed
            topics.append(topic)
            prefix = "Гипотеза" if (is_hypotheses_file or is_brainstorm) else "Пункт"
            enriched_lines.append(f"## {prefix} {num}: {topic}")
        else:
            enriched_lines.append(line)

    keywords: list[str] = []
    for topic in topics:
        keywords.extend(re.findall(r"[а-яёa-z]{5,}", topic.lower()))

    unique_kw: list[str] = []
    seen: set[str] = set()
    stop = {"изменение", "замена", "полная", "после", "провести", "опробование"}
    for kw in keywords:
        if kw not in seen and kw not in stop:
            seen.add(kw)
            unique_kw.append(kw)

    title = "Материалы предприятия" if is_hypotheses_file else "Структурированный документ"
    header = [f"# {title}: {path.name}"]
    if is_brainstorm or is_hypotheses_file:
        header.append("Нумерованные направления — приоритет для RAG.")
    if unique_kw:
        header.append("Ключевые темы: " + ", ".join(unique_kw[:12]))
    header.append("")
    return "\n".join(header + enriched_lines)


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        text = paragraph.text.strip()
        if not text:
            continue
        if "heading" in style:
            level = 2
            m = re.search(r"heading\s*(\d+)", style)
            if m:
                level = min(int(m.group(1)), 4)
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    for table_idx, table in enumerate(doc.tables, 1):
        parts.append(f"## Table {table_idx}")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return enrich_structured_docx("\n".join(parts), path)
