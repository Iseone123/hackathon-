"""Загрузка экспертных пар (Excel хвостов → DOCX гипотез) как few-shot примеров."""
from __future__ import annotations

import glob
import logging
from pathlib import Path

from docx import Document

from .tailings_parser import parse_tailings_xlsx

logger = logging.getLogger("few_shot")


def load_expert_pairs(data_dir: str) -> list[dict]:
    """Ищет пары Хвосты*.xlsx + Гипотезы*.docx в подпапках data_dir."""
    pairs = []
    for xlsx in sorted(glob.glob(f"{data_dir}/**/Хвосты*.xlsx", recursive=True)):
        folder = Path(xlsx).parent
        docx_files = list(folder.glob("Гипотезы*.docx"))
        if not docx_files:
            continue
        try:
            parsed = parse_tailings_xlsx(xlsx)
            hypotheses = read_hypotheses_docx(str(docx_files[0]))
        except Exception as e:
            logger.warning("Пара %s не загрузилась: %s", folder.name, e)
            continue
        pairs.append(
            {
                "name": Path(xlsx).stem.replace("Хвосты", "").strip(" _"),
                "xlsx": xlsx,
                "summary_text": parsed["summary_text"],
                "diagnostics": parsed["diagnostics"],
                "expert_hypotheses": hypotheses,
            }
        )
    logger.info("Загружено экспертных пар: %s", len(pairs))
    return pairs


def read_hypotheses_docx(path: str) -> list[str]:
    doc = Document(path)
    items: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in items:
                    items.append(text)
    for para in doc.paragraphs:
        t = para.text.strip()
        # строки вида "1. ..." вне таблиц
        if t and t[0].isdigit() and t not in items:
            items.append(t)
    return items
